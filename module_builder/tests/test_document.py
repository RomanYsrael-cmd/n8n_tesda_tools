from docx import Document

from app.docx_engine import build_module, unresolved_placeholders
from app.schemas import CourseMetadata
from app.validation import audit_docx, template_compatibility
from conftest import TEMPLATE


def test_template_has_every_required_placeholder():
    assert template_compatibility(TEMPLATE) == []


def test_build_preserves_tables_and_resolves_placeholders(tmp_path, bundle):
    source = Document(TEMPLATE)
    output = build_module(TEMPLATE, tmp_path, CourseMetadata(code="AE17", title="Accounting Information System", author="Trainer"), bundle)
    built = Document(output)
    assert len(built.tables) == len(source.tables)
    assert unresolved_placeholders(output) == []
    assert audit_docx(output, expected_tables=len(source.tables))["valid"]


def test_presentation_keep_with_next_applies_only_to_headings(tmp_path, bundle):
    output = build_module(TEMPLATE, tmp_path, CourseMetadata(code="AE17", title="Accounting Information System", author="Trainer"), bundle)
    paragraphs = Document(output).paragraphs
    heading = next(p for p in paragraphs if p.text == "Core concepts")
    body = paragraphs[paragraphs.index(heading) + 1]
    example = paragraphs[paragraphs.index(heading) + 2]

    assert heading.paragraph_format.keep_with_next is True
    assert body.paragraph_format.keep_with_next is False
    assert example.paragraph_format.keep_with_next is False
