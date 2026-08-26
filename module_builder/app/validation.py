from __future__ import annotations

import json
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from pydantic import ValidationError

from .docx_engine import PLACEHOLDERS, unresolved_placeholders
from .schemas import ModuleBundle, NormalizedSyllabus


def validate_bundle(raw: object, quiz_count: int) -> tuple[ModuleBundle | None, list[str]]:
    try:
        bundle = ModuleBundle.model_validate(raw)
    except ValidationError as exc:
        return None, [f"{'.'.join(map(str, e['loc']))}: {e['msg']}" for e in exc.errors()]
    errors = []
    if len(bundle.quiz.questions) != quiz_count:
        errors.append(f"quiz.questions: expected exactly {quiz_count}, received {len(bundle.quiz.questions)}")
    presentation = " ".join(bundle.presentation.presentation).casefold()
    for question in bundle.quiz.questions:
        keywords = [w.casefold().strip(".,?!:;()") for w in question.question.split() if len(w) > 5]
        if keywords and not any(k in presentation for k in keywords):
            errors.append(f"quiz question {question.id} may not be answerable from the presentation")
    return (bundle if not errors else None), errors


def validate_bundle_against_plan(bundle: ModuleBundle, plan: NormalizedSyllabus) -> list[str]:
    match = next((w for w in plan.weeks if w.actual_week == bundle.actual_week and w.generate), None)
    if not match:
        return [f"Week {bundle.actual_week} is not an approved generated week"]
    errors = []
    if match.lesson_number != bundle.lesson_number:
        errors.append(f"Week {bundle.actual_week} must be Lesson {match.lesson_number}")
    if bundle.approved_scope.strip().casefold() != match.topic_scope.strip().casefold():
        errors.append("approved_scope must exactly match the approved normalized plan")
    expected_title = f"Key Facts {bundle.lesson_number}.1 – {match.proposed_title}"
    if bundle.presentation.information_sheet_title.strip() != expected_title.strip():
        errors.append(f"information_sheet_title must be exactly: {expected_title}")
    return errors


def template_compatibility(path: Path) -> list[str]:
    doc = Document(path)
    from .docx_engine import iter_paragraphs
    text = "\n".join(p.text for p in iter_paragraphs(doc))
    return sorted(PLACEHOLDERS - {p for p in PLACEHOLDERS if p in text})


def audit_docx(path: Path, expected_images: int | None = None, expected_tables: int | None = None) -> dict:
    report = {"path": str(path), "valid": True, "errors": [], "warnings": []}
    try:
        doc = Document(path)
        if expected_tables is not None and len(doc.tables) != expected_tables:
            report["errors"].append(f"Expected {expected_tables} tables; found {len(doc.tables)}")
        with zipfile.ZipFile(path) as package:
            bad = package.testzip()
            images = len([n for n in package.namelist() if n.startswith("word/media/")])
            if bad:
                report["errors"].append(f"Corrupt package member: {bad}")
            if expected_images is not None and images != expected_images:
                report["errors"].append(f"Expected {expected_images} images; found {images}")
        report["errors"].extend(f"Unresolved placeholder: {x}" for x in unresolved_placeholders(path))
    except Exception as exc:
        report["errors"].append(f"DOCX could not be opened: {exc}")
    report["valid"] = not report["errors"]
    return report


def render_verify(path: Path, output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return {"valid": False, "skipped": True, "errors": ["LibreOffice is not installed in this environment"]}
    with tempfile.TemporaryDirectory(prefix="module-builder-lo-") as profile:
        result = subprocess.run([soffice, "--headless", f"-env:UserInstallation=file:///{Path(profile).as_posix()}", "--convert-to", "pdf", "--outdir", str(output_dir), str(path)], capture_output=True, text=True, timeout=180)
    pdf = output_dir / (path.stem + ".pdf")
    errors = []
    if result.returncode or not pdf.exists() or pdf.stat().st_size == 0:
        errors.append((result.stderr or result.stdout or "LibreOffice conversion failed").strip())
    return {"valid": not errors, "skipped": False, "pdf": str(pdf) if pdf.exists() else None, "errors": errors}


def report_json(path: Path, report: dict):
    path.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
