from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document

from .schemas import CourseMetadata, NormalizedSyllabus, SourceTrace, WeekPlan

WEEK_RE = re.compile(r"\bweeks?\s*(\d+)\s*(?:[-–—]|to)?\s*(\d+)?\b", re.I)
EXAM_RE = re.compile(r"\b(exam|examination|assessment week)\b", re.I)
ORIENTATION_RE = re.compile(r"\borientation\b", re.I)
SESSION_RE = re.compile(r"\b(preliminary|midterm|pre-final|final)\s+session\b", re.I)
LO_RE = re.compile(r"\bLO\s*\d+\s*[.:]", re.I)


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\x0b", " ")).strip()


def parse_week_range(text: str) -> list[int]:
    match = WEEK_RE.search(clean(text))
    if not match:
        return []
    start, end = int(match.group(1)), int(match.group(2) or match.group(1))
    if end < start or end - start > 100:
        raise ValueError(f"Invalid week range: {text}")
    return list(range(start, end + 1))


def split_items(text: str) -> tuple[str, list[str]]:
    text = clean(text)
    parts = [clean(x) for x in re.split(r"\s+(?=\d+[.)]\s*)|[\n\r]+", text) if clean(x)]
    title = re.sub(r"\s*\d+[.)].*$", "", text).strip(" :-")
    details = [re.sub(r"^\d+[.)]\s*", "", p) for p in parts if re.match(r"^\d+[.)]", p)]
    return title or text, details


def split_scope(topic: str, weeks: list[int]) -> list[tuple[str, str]]:
    title, details = split_items(topic)
    if len(weeks) == 1:
        return [(title, topic)]
    if not details:
        return [(f"{title} - Part {i + 1}", f"{title}, Part {i + 1} of {len(weeks)}") for i in range(len(weeks))]
    groups: list[list[str]] = [[] for _ in weeks]
    for i, item in enumerate(details):
        groups[min(i * len(weeks) // len(details), len(weeks) - 1)].append(item)
    for i, group in enumerate(groups):
        if not group:
            group.append(f"Integrated application of {title} concepts from earlier parts")
    return [
        (f"{title} - Part {i + 1}: {group[0]}", f"{title}: " + "; ".join(group))
        for i, group in enumerate(groups)
    ]


def _cell_text(cell) -> str:
    seen = []
    for p in cell.paragraphs:
        value = clean(p.text)
        if value and value not in seen:
            seen.append(value)
    return " | ".join(seen)


def extract_syllabus(path: Path) -> NormalizedSyllabus:
    doc = Document(path)
    code = ""
    title = path.stem
    description = ""
    outcomes: list[str] = []
    references: list[str] = []
    author = ""

    for table in doc.tables:
        for row in table.rows:
            cells = [_cell_text(c) for c in row.cells]
            label = cells[0].upper() if cells else ""
            values = [x for x in cells[1:] if x and x != ":"]
            if "COURSE CODE" in label and values:
                code = values[0]
            elif "COURSE TITLE" in label and values:
                title = values[0]
            elif "COURSE DESCRIPTION" in label and values:
                description = values[0]
            for value in cells:
                if re.match(r"^CO\d+[.:]", value, re.I) and value not in outcomes:
                    outcomes.append(value)
    all_paras = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    for i, text in enumerate(all_paras):
        if text.lower().startswith("prepared by") and i + 1 < len(all_paras):
            author = all_paras[i + 1]
        if "reference" in text.lower() and i + 1 < len(all_paras):
            references.extend(x for x in all_paras[i + 1 :] if len(x) > 20 and not x.lower().startswith(("prepared", "noted", "reviewed", "approved")))

    week_plans: list[WeekPlan] = []
    current_session = ""
    current_outcome = ""
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = [_cell_text(c) for c in row.cells]
            joined = " | ".join(dict.fromkeys(x for x in cells if x))
            session_match = SESSION_RE.search(joined)
            if session_match and not WEEK_RE.search(joined):
                current_session = session_match.group(0)
            if LO_RE.search(cells[0] if cells else "") and not WEEK_RE.search(joined):
                current_outcome = cells[0]
            week_cell = next((x for x in reversed(cells) if WEEK_RE.search(x)), "")
            weeks = parse_week_range(week_cell) if week_cell else []
            if not weeks:
                continue
            topic = cells[0] if cells else "Untitled topic"
            kind = "examination" if EXAM_RE.search(topic) else "orientation" if ORIENTATION_RE.search(topic) else "instruction"
            generate = kind == "instruction"
            split = split_scope(topic, weeks)
            for index, week in enumerate(weeks):
                proposed, scope = split[index]
                trace = [SourceTrace(table_index=ti, row_index=ri, cell_index=ci, source_text=value) for ci, value in enumerate(cells) if value]
                week_plans.append(WeekPlan(
                    actual_week=week,
                    proposed_title=proposed,
                    learning_outcome=current_outcome,
                    generate=generate,
                    type=kind,
                    session=current_session,
                    topic_scope=scope,
                    methods=[x for x in re.split(r"\s*\|\s*", cells[1])] if len(cells) > 1 and cells[1] else [],
                    presentation_guidance=cells[2] if len(cells) > 2 else "",
                    practice=cells[3] if len(cells) > 3 else "",
                    feedback=cells[4] if len(cells) > 4 else "",
                    resources=[x for x in re.split(r"\s*\|\s*", cells[6])] if len(cells) > 6 and cells[6] else [],
                    skipped_reason="Skipped by default: orientation" if kind == "orientation" else "Skipped by default: examination" if kind == "examination" else "",
                    multi_week_source=week_cell if len(weeks) > 1 else "",
                    warnings=["Review automatically split multi-week scope"] if len(weeks) > 1 else [],
                    source_traceability=trace,
                ))
    by_week = {w.actual_week: w for w in week_plans}
    week_plans = [by_week[k] for k in sorted(by_week)]
    lesson = 0
    for week in week_plans:
        if week.generate:
            lesson += 1
            week.lesson_number = lesson
        else:
            week.lesson_number = None
    warnings = []
    if not week_plans:
        warnings.append("No week schedule was detected; review the source document layout")
    return NormalizedSyllabus(
        source_filename=path.name,
        course=CourseMetadata(code=code, title=title, description=description, author=author, outcomes=outcomes, references=references),
        weeks=week_plans,
        normalization_warnings=warnings,
    )


def renumber(weeks: list[WeekPlan]) -> list[WeekPlan]:
    lesson = 0
    for week in sorted(weeks, key=lambda w: w.actual_week):
        if week.generate:
            lesson += 1
            week.lesson_number = lesson
            week.skipped_reason = ""
        else:
            week.lesson_number = None
    return weeks

