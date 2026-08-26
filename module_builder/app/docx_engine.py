from __future__ import annotations

import copy
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from .schemas import CourseMetadata, ModuleBundle, PresentationBlock, RichTextSpan
from .storage import safe_module_filename

PLACEHOLDERS = {
    "{{subject}}", "{{author}}", "{{trainer}}", "{{weekX}}", "{{lessonX}}", "{{lessonXtitle}}",
    "{{list_of_LO}}", "{{preassessment}}", "{{presentation}}", "{{letsexercise}}", "{{contents_mc}}",
    "{{LE_answer_key}}", "{{la_title}}", "{{la_objective}}", "{{la_sup_mat}}", "{{la_equipment_list}}",
    "{{la_steps_list}}", "{{la_assessmentmethod}}", "{{la_pc1}}", "{{la_pc2}}", "{{la_pc3}}", "{{la_pc4}}", "{{la_pc5}}",
}


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            yield from ncell.paragraphs


def _replace_text(paragraph: Paragraph, replacements: dict[str, str]):
    text = paragraph.text
    changed = text
    for key, value in replacements.items():
        changed = changed.replace(key, value)
    if changed == text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = changed
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(changed)


def _insert_after(paragraph: Paragraph, text: str, numbered: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    if paragraph._p.pPr is not None:
        new_p.append(copy.deepcopy(paragraph._p.pPr))
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if numbered:
        ppr = result._p.get_or_add_pPr()
        numpr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0")
        numid = OxmlElement("w:numId"); numid.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "5")
        numpr.extend([ilvl, numid]); ppr.append(numpr)
    result.add_run(text)
    return result


def _fill_list(paragraph: Paragraph, marker: str, items: list[str], numbered: bool = False):
    if paragraph.text.strip() != marker:
        _replace_text(paragraph, {marker: "\n".join(items)})
        return
    if not paragraph.runs:
        paragraph.add_run()
    paragraph.runs[0].text = items[0] if items else ""
    current = paragraph
    for item in items[1:]:
        current = _insert_after(current, item, numbered=numbered)


def _set_rich_block(paragraph: Paragraph, block: PresentationBlock, previous_heading: str = ""):
    for run in paragraph.runs:
        run.text = ""
    # Inserted paragraphs inherit the preceding block's paragraph properties.
    # Make pagination explicit so a heading stays with its first body paragraph
    # without accidentally chaining the entire information sheet together.
    paragraph.paragraph_format.keep_with_next = block.type == "heading"
    spans = list(block.spans)
    if previous_heading and block.type in {"paragraph", "bullet", "numbered"} and spans:
        if spans[0].text.strip().rstrip(":").casefold() == previous_heading.strip().rstrip(":").casefold():
            spans = spans[1:]
    if spans and block.type in {"example", "note"}:
        prefix_pattern = r"^\s*(?:example\s*:\s*(?:realistic\s+example\s*:\s*)?|realistic\s+example\s*:\s*|note\s*:\s*)"
        cleaned = re.sub(prefix_pattern, "", spans[0].text, flags=re.IGNORECASE)
        spans[0] = RichTextSpan(text=cleaned or " ", bold=spans[0].bold, italic=spans[0].italic)
    if block.type == "example":
        prefix = paragraph.add_run("Example: ")
        prefix.bold = True
    elif block.type == "note":
        prefix = paragraph.add_run("Note: ")
        prefix.bold = True
    for span in spans:
        run = paragraph.add_run(span.text)
        run.bold = span.bold or block.type == "heading"
        run.italic = span.italic or block.type == "note"
    if block.type in {"bullet", "numbered"}:
        try:
            paragraph.style = "List Bullet" if block.type == "bullet" else "List Number"
        except KeyError:
            pass


def _fill_rich_blocks(paragraph: Paragraph, blocks: list[PresentationBlock]):
    current = paragraph
    previous_heading = ""
    for index, block in enumerate(blocks):
        if index:
            current = _insert_after(current, "")
        _set_rich_block(current, block, previous_heading)
        previous_heading = block.plain_text if block.type == "heading" else ""


def build_module(template: Path, output_dir: Path, course: CourseMetadata, bundle: ModuleBundle) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    title = bundle.presentation.lesson_title
    output = output_dir / safe_module_filename(bundle.actual_week, bundle.lesson_number, title)
    shutil.copy2(template, output)
    doc = Document(output)
    qlines = []
    for index, question in enumerate(bundle.quiz.questions, 1):
        choices = "\n".join(f"   {c.id}. {c.text}" for c in question.choices)
        qlines.append(f"{index}. {question.question}\n{choices}")
    answer_lines = [f"{i + 1}. {bundle.quiz.answer_key[q.id]}" for i, q in enumerate(bundle.quiz.questions)]
    replacements = {
        "{{subject}}": f"{course.code} {course.title}".strip(), "{{author}}": course.author,
        "{{trainer}}": course.trainer or course.author, "{{weekX}}": str(bundle.actual_week),
        "{{lessonX}}": str(bundle.lesson_number), "{{lessonXtitle}}": title,
        "{{letsexercise}}": "Let's Exercise", "{{la_title}}": bundle.practical_activity.title,
        "{{la_objective}}": bundle.practical_activity.performance_objective,
        "{{la_assessmentmethod}}": bundle.practical_activity.assessment_method,
    }
    for i, criterion in enumerate(bundle.practical_activity.performance_criteria, 1):
        replacements[f"{{{{la_pc{i}}}}}"] = criterion
    reference_blocks = []
    if bundle.presentation.references:
        reference_blocks.append(PresentationBlock(type="heading", spans=[RichTextSpan(text="References")]))
        for ref in bundle.presentation.references:
            details = ". ".join(part for part in [ref.author_or_organization, ref.year, ref.title, ref.url] if part)
            reference_blocks.append(PresentationBlock(type="bullet", spans=[RichTextSpan(text=details)]))
    presentation_blocks = [PresentationBlock(type="heading", spans=[RichTextSpan(text=bundle.presentation.information_sheet_title)]), PresentationBlock(type="paragraph", spans=[RichTextSpan(text=bundle.presentation.introduction)]), *bundle.presentation.presentation, *reference_blocks]
    list_slots = {
        "{{list_of_LO}}": bundle.presentation.measurable_objectives,
        "{{preassessment}}": bundle.presentation.pre_assessment,
        "{{contents_mc}}": qlines,
        "{{LE_answer_key}}": answer_lines,
        "{{la_sup_mat}}": bundle.practical_activity.supplies_materials or ["None required"],
        "{{la_equipment_list}}": bundle.practical_activity.equipment or ["None required"],
        "{{la_steps_list}}": bundle.practical_activity.steps,
    }
    for paragraph in list(iter_paragraphs(doc)):
        if "{{presentation}}" in paragraph.text:
            _fill_rich_blocks(paragraph, presentation_blocks)
            continue
        for marker, items in list_slots.items():
            if marker in paragraph.text:
                _fill_list(paragraph, marker, items, numbered=marker in {"{{la_steps_list}}"})
                break
        _replace_text(paragraph, replacements)
    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            run.font.name = course.font_family
            run.font.size = Pt(course.font_size)
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), course.font_family)
    doc.save(output)
    return output


def unresolved_placeholders(path: Path) -> list[str]:
    doc = Document(path)
    text = "\n".join(p.text for p in iter_paragraphs(doc))
    return sorted(set(re.findall(r"\{\{[^{}]+\}\}", text)))


def package_course(course_dir: Path, zip_name: str = "course-modules.zip") -> Path:
    output = course_dir / zip_name
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in course_dir.rglob("*"):
            if path.is_file() and path != output:
                archive.write(path, path.relative_to(course_dir))
    return output
