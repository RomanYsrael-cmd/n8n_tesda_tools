from pathlib import Path

from cblm_app.documents import audit_docx, build_cblm, iter_paragraphs
from docx import Document
from cblm_app.parsers import parse_self_check, parse_task_sheet
from cblm_app.prompts import PromptCatalog
from cblm_app.schemas import CBLMCourse, CBLMLearningOutcome, CBLMPlan, CBLMTopic

ROOT = Path(__file__).parents[1]


def test_prompt_catalog_substitutes_full_values():
    text = PromptCatalog.load(ROOT / "Prompts.xlsx").render("module_title", {"learning_outcome": "Perform testing"})
    assert "Perform testing" in text and "{{learning_outcome}}" not in text


def test_python_derives_validated_sections():
    quiz = parse_self_check("Questions\n1. Test?\nA. One\nB. Two\nAnswer Key\n1. A")
    assert quiz.quiz and quiz.answer_key
    task = parse_task_sheet("""Title: Build an output
Performance Objective: Produce it
List of Supplies: Paper
List of Equipment: None
Steps: 1. Build it
Assessment Method: Direct observation
Performance Criteria:
- Output is complete
- Output is accurate""")
    assert task.activity_title == "Build an output" and len(task.activity_criteria) == 2


def test_task_sheet_infers_unlabelled_title_and_cleans_markdown_checklist():
    task = parse_task_sheet("""TASK SHEET
**Analyzing Workplace Identity Scenarios**

**Performance Objective**
Analyze three cases.
**Supplies/Materials**
• Worksheet
**Equipment**
• None required
**Steps/Procedure**
1. Analyze each case.
**Assessment Method**
Work Sample / Output and Direct Observation
**Performance Criteria**
| Performance Criteria | Yes | No |
| :--- | :---: | :---: |
| 1. The output identifies the perspective. | ☐ | ☐ |
| 2. The intervention is practical. | ☐ | ☐ |
**Critical Criteria**
This explanatory metadata must not become a criterion.
**Task Difficulty**
Intermediate""")
    assert task.activity_title == "Analyzing Workplace Identity Scenarios"
    assert task.activity_criteria == ["The output identifies the perspective.", "The intervention is practical."]


def test_builds_complete_markdown_aware_docx(tmp_path):
    topic = CBLMTopic(number=1,title="Safety",weeks=[1],learning_objectives="## Objectives\n- Explain **safety**",keyfacts_content="# Safety\nUse *care*.\n\n---\n\n### Comparison\n| Safe | Unsafe |\n| :--- | :--- |\n| PPE | No PPE |\n\n- Main item\n  - Nested item",quiz_instructions="Choose.",quiz="1. Safe?\nA. Yes\nB. No",answer_key="1. A",activity_title="Inspection",activity_objectives="Inspect",activity_supplies="- Checklist",activity_equipment="- PPE",activity_steps="1. Inspect",activity_method="Direct observation",activity_criteria=["Checks hazards"])
    lo = CBLMLearningOutcome(number=1,learning_outcome="Apply safety",duration=3,topics=[topic])
    plan = CBLMPlan(source_filename="sample.docx",course=CBLMCourse(course_title="Test Course"),learning_outcomes=[lo])
    path = build_cblm(ROOT / "Templates", tmp_path, plan, lo)
    report = audit_docx(path)
    assert path.exists() and report["valid"], report
    doc = Document(path)
    assert any(table.cell(0, 0).text == "Safe" and table.cell(1, 1).text == "No PPE" for table in doc.tables if len(table.rows) > 1 and len(table.columns) > 1)
    assert not any(paragraph.text.strip() == "---" for paragraph in doc.paragraphs)
    assert any(paragraph.text == "Comparison" and paragraph.style.name == "Heading 3" for paragraph in doc.paragraphs)
    nested = next(paragraph for paragraph in iter_paragraphs(doc) if paragraph.text == "Nested item")
    assert nested.paragraph_format.left_indent and nested.paragraph_format.left_indent.inches >= 0.5


def test_lists_each_topic_merges_instructions_and_applies_font(tmp_path):
    topics = [
        CBLMTopic(number=1,title="First topic",weeks=[1],learning_objectives="Learn one",keyfacts_content="Facts",quiz_instructions="Choose",quiz="1. Q",answer_key="1. A",activity_title="First task",activity_objectives="Do",activity_supplies="Paper",activity_equipment="None",activity_steps="1. Do",activity_method="Direct observation",activity_criteria=["First output is correct"],resources=["Book"]),
        CBLMTopic(number=2,title="Second topic",weeks=[2],learning_objectives="Learn two",keyfacts_content="Facts",quiz_instructions="Choose",quiz="1. Q",answer_key="1. A",activity_title="Second task",activity_objectives="Do",activity_supplies="Paper",activity_equipment="None",activity_steps="1. Do",activity_method="Written test/quiz",activity_criteria=["Second output is correct"],resources=["LMS"]),
    ]
    lo = CBLMLearningOutcome(number=1,learning_outcome="Perform work",duration=6,training_materials=["Book","LMS"],topics=topics)
    plan = CBLMPlan(source_filename="sample.docx",course=CBLMCourse(course_title="Course",font_family="Bookman Old Style",font_size=12),learning_outcomes=[lo])
    path = build_cblm(ROOT/"Templates",tmp_path,plan,lo); doc=Document(path)
    text="\n".join(p.text for p in doc.paragraphs)+"\n"+"\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "First topic" in text and "Second topic" in text and "{{" not in text
    assert "First output is correct" in text and "Second output is correct" in text
    assert all(run.font.name == "Bookman Old Style" for p in doc.paragraphs for run in p.runs if run.text)
