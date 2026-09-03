from __future__ import annotations

import copy
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from .schemas import CBLMLearningOutcome, CBLMPlan, CBLMTopic
from .storage import safe_name


TOKEN_RE = re.compile(r"\{\{[^{}]+\}\}")
INLINE_RE = re.compile(r"(\*\*\*.+?\*\*\*|___[^_]+___|\*\*.+?\*\*|__.+?__|(?<!\*)\*[^*]+\*(?!\*)|(?<!_)_[^_]+_(?!_))")


def iter_paragraphs(doc: Document):
    containers = [doc, *(section.header for section in doc.sections), *(section.footer for section in doc.sections)]
    def cell_paragraphs(cell):
        yield from cell.paragraphs
        for nested_table in cell.tables:
            for nested_row in nested_table.rows:
                for nested_cell in nested_row.cells:
                    yield from cell_paragraphs(nested_cell)
    for container in containers:
        yield from container.paragraphs
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell_paragraphs(cell)


def inline_parts(text: str):
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            yield text[position:match.start()], False, False
        token = match.group(0)
        both = token.startswith(("***", "___"))
        bold = both or token.startswith(("**", "__"))
        italic = both or (not bold)
        trim = 3 if both else 2 if bold else 1
        yield token[trim:-trim], bold, italic
        position = match.end()
    if position < len(text):
        yield text[position:], False, False


def _copy_run_format(source, target):
    if source._r.rPr is not None:
        target._r.insert(0, copy.deepcopy(source._r.rPr))


def replace_inline(paragraph: Paragraph, marker: str, value: str):
    if paragraph.text.count(marker) > 1:
        full_text = paragraph.text.replace(marker, value)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(full_text)
        return True
    for run in list(paragraph.runs):
        if marker not in run.text:
            continue
        before, after = run.text.split(marker, 1)
        run.text = before
        cursor = run
        for text, bold, italic in inline_parts(value.strip()):
            new_r = OxmlElement("w:r")
            cursor._r.addnext(new_r)
            new_run = type(run)(new_r, paragraph)
            _copy_run_format(run, new_run)
            new_run.text = text
            new_run.bold = bold or new_run.bold
            new_run.italic = italic or new_run.italic
            cursor = new_run
        if after:
            new_r = OxmlElement("w:r")
            cursor._r.addnext(new_r)
            new_run = type(run)(new_r, paragraph)
            _copy_run_format(run, new_run)
            new_run.text = after
        return True
    # Word may split a visible placeholder across several runs after a user
    # edits the template. Preserve the paragraph's leading formatting while
    # deterministically consolidating and replacing the visible text.
    if marker in paragraph.text:
        full_text = paragraph.text.replace(marker, value)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(full_text)
        return True
    return False


def _new_after(paragraph: Paragraph) -> Paragraph:
    element = OxmlElement("w:p")
    if paragraph._p.pPr is not None:
        element.append(copy.deepcopy(paragraph._p.pPr))
    paragraph._p.addnext(element)
    return Paragraph(element, paragraph._parent)


def _set_paragraph_md(paragraph: Paragraph, text: str, style: str = ""):
    for run in paragraph.runs:
        run.text = ""
    if style:
        try:
            paragraph.style = style
        except KeyError:
            pass
    for value, bold, italic in inline_parts(text):
        run = paragraph.add_run(value)
        run.bold = bold
        run.italic = italic


def markdown_blocks(value: str):
    lines = value.strip().strip("`").splitlines()
    paragraph = []
    blocks = []
    def flush():
        if paragraph:
            blocks.append(("paragraph", " ".join(paragraph), 0))
            paragraph.clear()
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush(); index += 1; continue
        if re.fullmatch(r"(?:-{3,}|\*{3,}|_{3,})", line):
            flush(); blocks.append(("rule", "", 0)); index += 1; continue
        if line.startswith("|") and line.endswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1]):
            flush(); rows = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells): rows.append(cells)
                index += 1
            blocks.append(("table", rows, 0)); continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        bullet = re.match(r"^(\s*)[-*+]\s+(.+)$", raw)
        numbered = re.match(r"^(\s*)\d+[.)]\s+(.+)$", raw)
        if heading:
            flush(); blocks.append(("heading", heading.group(2), len(heading.group(1))))
        elif bullet:
            flush(); blocks.append(("bullet", bullet.group(2), len(bullet.group(1).expandtabs(4)) // 2))
        elif numbered:
            flush(); blocks.append(("numbered", numbered.group(2), len(numbered.group(1).expandtabs(4)) // 2))
        elif line.startswith(">"):
            flush(); blocks.append(("quote", line.lstrip("> "), 0))
        else:
            paragraph.append(line.replace("[ ]", "☐").replace("[x]", "☒").replace("[X]", "☒"))
        index += 1
    flush()
    return blocks or [("paragraph", "", 0)]


def _paragraph_after_element(element, parent) -> Paragraph:
    xml = OxmlElement("w:p")
    element.addnext(xml)
    return Paragraph(xml, parent)


def _set_rule(paragraph: Paragraph):
    paragraph.text = ""
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for key, val in (("val", "single"), ("sz", "6"), ("space", "4"), ("color", "808080")):
        bottom.set(qn(f"w:{key}"), val)
    borders.append(bottom); ppr.append(borders)


def _table_after(element, paragraph: Paragraph, rows: list[list[str]]):
    columns = max((len(row) for row in rows), default=1)
    table = paragraph.part.document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    for rindex, row in enumerate(rows):
        for cindex, value in enumerate(row):
            cell_p = table.cell(rindex, cindex).paragraphs[0]
            _set_paragraph_md(cell_p, value)
            if rindex == 0:
                for run in cell_p.runs: run.bold = True
    element.addnext(table._tbl)
    return table._tbl


def replace_rich(paragraph: Paragraph, marker: str, value: str):
    if marker not in paragraph.text:
        return False
    if paragraph.text.strip() != marker:
        return replace_inline(paragraph, marker, re.sub(r"[*_`]", "", value))
    current_element = paragraph._p
    used_anchor = False
    for kind, text, level in markdown_blocks(value):
        if kind == "table":
            current_element = _table_after(current_element, paragraph, text)
            continue
        current = paragraph if not used_anchor else _paragraph_after_element(current_element, paragraph._parent)
        used_anchor = True
        current_element = current._p
        if kind == "rule":
            _set_rule(current); continue
        if kind == "heading":
            style = f"Heading {min(level, 3)}"
        elif kind == "bullet":
            style = "List Bullet" if level == 0 else "List Bullet 2"
        elif kind == "numbered":
            style = "List Number" if level == 0 else "List Number 2"
        elif kind == "quote":
            style = "Quote"
        else:
            style = ""
        _set_paragraph_md(current, text, style)
        if kind in {"bullet", "numbered"} and level:
            current.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        current.paragraph_format.keep_with_next = kind == "heading"
    if not used_anchor and paragraph.text:
        paragraph.text = ""
    return True


def clone_paragraph_items(paragraph: Paragraph, marker: str, items: list[str]):
    values = items or [""]
    template_xml = copy.deepcopy(paragraph._p)
    replace_rich(paragraph, marker, values[0])
    current = paragraph
    for value in values[1:]:
        new_p = copy.deepcopy(template_xml)
        current._p.addnext(new_p)
        current = Paragraph(new_p, paragraph._parent)
        replace_rich(current, marker, value)


def replace_all(doc: Document, replacements: dict[str, str], rich: dict[str, str] | None = None):
    rich = rich or {}
    for paragraph in list(iter_paragraphs(doc)):
        for marker, value in rich.items():
            if marker in paragraph.text:
                replace_rich(paragraph, marker, value)
        for marker, value in replacements.items():
            if marker in paragraph.text:
                replace_inline(paragraph, marker, value)


def _clone_row_after(row):
    clone = copy.deepcopy(row._tr)
    row._tr.addnext(clone)
    return clone


def _fill_competency_table(doc: Document, outcomes: list[CBLMLearningOutcome]):
    table = doc.tables[1]
    template = table.rows[1]
    rows = [template]
    for _ in outcomes[1:]:
        _clone_row_after(rows[-1])
        rows = table.rows[1:]
    for index, (row, outcome) in enumerate(zip(table.rows[1:], outcomes), 1):
        row.cells[0].text = str(index)
        replace_all_in_cell(row.cells[1], {"{{learning_outcome}}": outcome.learning_outcome})
        replace_all_in_cell(row.cells[2], {"{{module_title}}": outcome.module_title})


def replace_all_in_cell(cell, replacements):
    paragraphs = list(cell.paragraphs)
    for table in cell.tables:
        for row in table.rows:
            for nested in row.cells:
                paragraphs.extend(nested.paragraphs)
    for paragraph in paragraphs:
        for marker, value in replacements.items():
            replace_inline(paragraph, marker, value)


def unique(values):
    return list(dict.fromkeys(str(value).strip() for value in values if str(value).strip()))


def _fill_learning_experience(doc: Document, outcome: CBLMLearningOutcome):
    table = doc.tables[0]
    base_rows = [copy.deepcopy(table.rows[index]._tr) for index in range(1, 4)]
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    for topic_index, topic in enumerate(outcome.topics, 1):
        for row_xml in base_rows:
            clone = copy.deepcopy(row_xml)
            table._tbl.append(clone)
        rows = table.rows[-3:]
        mapping = {"{{x}}": str(outcome.number), "{{y}}": str(topic_index), "{{syllabus_topic}}": topic.title, "{{activity_title}}": topic.activity_title}
        for row in rows:
            replace_all_in_cell(row.cells[0], mapping)
            if topic_index > 1:
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.text = ""
    # The instructions apply to the whole learning outcome. Make the right
    # side one continuous vertically merged cell across every activity row.
    right_text = table.cell(1, 1).text
    for row in table.rows[1:]:
        row.cells[1].text = ""
    merged = table.cell(1, 1).merge(table.cell(len(table.rows) - 1, 1))
    merged.text = right_text


def _fill_activity_criteria(doc: Document, criteria: list[str]):
    table = doc.tables[1]
    template = table.rows[1]
    for _ in criteria[1:]:
        _clone_row_after(table.rows[-1])
    for row, criterion in zip(table.rows[1:], criteria or [""]):
        replace_all_in_cell(row.cells[0], {"{{activity_criteria}}": criterion})


def _base_values(plan: CBLMPlan, outcome: CBLMLearningOutcome):
    return {
        "{{sector}}": plan.course.sector,
        "{{course_title}}": plan.course.course_title,
        "{{course_code}}": plan.course.course_code,
        "{{module_title}}": outcome.module_title,
        "{{name}}": plan.course.name,
        "{{learning_outcome}}": outcome.learning_outcome,
        "{{x}}": str(outcome.number),
        "{{next_learning_outcome}}": outcome.next_learning_outcome or "End of course",
        "{{module_descriptor}}": outcome.module_descriptor,
        # The legacy template still contains this removed field. Emptying it
        # preserves the surrounding design without introducing another call.
        "{{introduction}}": "",
        "{{duration}}": f"{outcome.duration:g} Hours",
        "{{location}}": outcome.laboratory or outcome.location,
    }


def apply_uniform_font(doc: Document, family: str, size: float):
    for style in doc.styles:
        if style.type == 1:
            style.font.name = family
            style.font.size = Pt(size)
            style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), family)
    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            run.font.name = family
            run.font.size = Pt(size)
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), family)


def _fragment(template_dir: Path, name: str, plan: CBLMPlan, outcome: CBLMLearningOutcome, topic: CBLMTopic | None = None) -> Document:
    doc = Document(template_dir / name)
    values = _base_values(plan, outcome)
    if name == "00_front_matter.docx":
        _fill_competency_table(doc, plan.learning_outcomes)
        lo_anchor = next(p for p in doc.paragraphs if "{{learning_outcome_specific}}" in p.text)
        clone_paragraph_items(lo_anchor, "{{learning_outcome_specific}}", [item.title for item in outcome.topics])
        criteria = unique(criterion for item in outcome.topics for criterion in item.activity_criteria)
        ac_anchor = next(p for p in doc.paragraphs if "{{assessment_criteria}}" in p.text)
        clone_paragraph_items(ac_anchor, "{{assessment_criteria}}", criteria or ["To be generated from the approved learning activities"])
    elif name == "10_lo_intro.docx":
        topics_marker = "{{learning_outcome_specific}}" if any("{{learning_outcome_specific}}" in p.text for p in doc.paragraphs) else "{{syllabus_topic}}"
        topics_anchor = next(p for p in doc.paragraphs if topics_marker in p.text)
        clone_paragraph_items(topics_anchor, topics_marker, [item.title for item in outcome.topics])
        criteria = unique(criterion for item in outcome.topics for criterion in item.activity_criteria)
        criteria_anchor = next(p for p in doc.paragraphs if "{{assessment_criteria}}" in p.text)
        clone_paragraph_items(criteria_anchor, "{{assessment_criteria}}", criteria or ["To be demonstrated through the learning activities"])
        materials = unique(outcome.training_materials or [resource for item in outcome.topics for resource in item.resources]) or ["Writing materials"]
        methods = unique(item.activity_method for item in outcome.topics) or ["Direct observation"]
        clone_paragraph_items(next(p for p in doc.paragraphs if "{{training_materials}}" in p.text), "{{training_materials}}", materials)
        clone_paragraph_items(next(p for p in doc.paragraphs if "{{assessment_methods}}" in p.text), "{{assessment_methods}}", methods)
    elif name == "20_learning_experience.docx":
        _fill_learning_experience(doc, outcome)
    elif topic is not None:
        values.update({"{{y}}": str(topic.number), "{{syllabus_topic}}": topic.title})
        if name == "30_key_facts.docx":
            replace_all(doc, values, {"{{lesson_objectives}}": topic.learning_objectives, "{{keyfacts_content}}": topic.keyfacts_content})
            return doc
        if name == "40_lets_exercise.docx":
            replace_all(doc, values, {"{{quiz_instructions}}": topic.quiz_instructions, "{{quiz}}": topic.quiz, "{{answer_key}}": topic.answer_key})
            return doc
        if name == "50_lets_apply.docx":
            _fill_activity_criteria(doc, topic.activity_criteria)
            values.update({
                "{{activity_title}}": topic.activity_title,
                "{{activity_objectives}}": topic.activity_objectives,
                "{{activity_supplies}}": topic.activity_supplies,
                "{{activity_equipment}}": topic.activity_equipment,
                "{{activity_steps}}": topic.activity_steps,
                "{{activity_method}}": topic.activity_method,
            })
            replace_all(doc, values)
            return doc
    replace_all(doc, values)
    return doc


def _append_document(master: Document, fragment: Document):
    if master.paragraphs:
        master.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    for element in fragment.element.body:
        if element.tag.endswith("}sectPr"):
            continue
        master.element.body.insert(-1, copy.deepcopy(element))


def build_cblm(template_dir: Path, output_dir: Path, plan: CBLMPlan, outcome: CBLMLearningOutcome) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    master = _fragment(template_dir, "00_front_matter.docx", plan, outcome)
    _append_document(master, _fragment(template_dir, "10_lo_intro.docx", plan, outcome))
    _append_document(master, _fragment(template_dir, "20_learning_experience.docx", plan, outcome))
    for topic in outcome.topics:
        _append_document(master, _fragment(template_dir, "30_key_facts.docx", plan, outcome, topic))
        _append_document(master, _fragment(template_dir, "40_lets_exercise.docx", plan, outcome, topic))
        _append_document(master, _fragment(template_dir, "50_lets_apply.docx", plan, outcome, topic))
    apply_uniform_font(master, plan.course.font_family, plan.course.font_size)
    filename = safe_name(f"CBLM {outcome.number:02d} - {outcome.module_title or outcome.learning_outcome}")
    path = output_dir / filename
    master.save(path)
    return path


def audit_docx(path: Path):
    try:
        doc = Document(path)
        unresolved = sorted(set(TOKEN_RE.findall("\n".join(p.text for p in iter_paragraphs(doc)))))
        with zipfile.ZipFile(path) as package:
            bad = package.testzip()
        errors = (["Unresolved placeholders: " + ", ".join(unresolved)] if unresolved else []) + ([f"Invalid DOCX package member: {bad}"] if bad else [])
        return {"valid": not errors, "errors": errors, "pages_checked_by_render": False}
    except Exception as exc:
        return {"valid": False, "errors": [str(exc)], "pages_checked_by_render": False}


def package_outputs(base: Path) -> Path:
    output = base / "cblm-course.zip"
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in base.rglob("*"):
            if path.is_file() and path != output and "render" not in path.parts:
                archive.write(path, path.relative_to(base))
    return output
